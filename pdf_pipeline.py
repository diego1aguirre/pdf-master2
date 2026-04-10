"""
PDF pipeline: convert DOCX → PDF, merge PDFs in order, optionally add page numbers.
All steps are separate, pure functions.
"""

import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import List

from pypdf import PdfReader, PdfWriter


def convert_docx_to_pdf(docx_path: Path, output_dir: Path | None = None) -> Path:
    """
    Convert a single .docx file to PDF.

    On macOS we prefer LibreOffice (soffice) because docx2pdf/JXA + Word often
    fails with 'Error: Message not understood'. On Windows we prefer docx2pdf
    (Word automation) and fall back to LibreOffice if available.

    Returns the path to the generated PDF.
    """
    docx_path = Path(docx_path).resolve()
    if not docx_path.suffix.lower() == ".docx":
        raise ValueError(f"Not a .docx file: {docx_path}")
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    out_dir = Path(output_dir) if output_dir else docx_path.parent
    out_dir = out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / f"{docx_path.stem}.pdf"

    def _try_docx2pdf() -> bool:
        try:
            from docx2pdf import convert as docx2pdf_convert

            docx2pdf_convert(str(docx_path), str(pdf_path))
            return pdf_path.exists()
        except Exception:
            return False

    def _try_soffice() -> bool:
        # LibreOffice headless: soffice --headless --convert-to pdf --outdir out_dir input.docx
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice and Path("/Applications/LibreOffice.app").exists():
            soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if not soffice or not Path(soffice).exists():
            return False
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
                check=True,
                capture_output=True,
                timeout=300,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            return False
        return pdf_path.exists()

    def _try_mammoth_weasyprint() -> bool:
        # Pure-Python fallback with formatting: mammoth converts docx→HTML
        # (preserving bold, italic, headings, tables, lists, images), then
        # weasyprint renders the HTML to PDF.
        try:
            import base64
            import mammoth
            import weasyprint

            # Embed images as base64 data URIs so they appear in the PDF
            def _inline_image(image):
                with image.open() as img_bytes:
                    encoded = base64.b64encode(img_bytes.read()).decode("ascii")
                return {"src": f"data:{image.content_type};base64,{encoded}"}

            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_html(
                    f,
                    convert_image=mammoth.images.img_element(_inline_image),
                )

            html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{ margin: 2cm; }}
  body {{ font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #000; }}
  h1 {{ font-size: 20pt; font-weight: bold; margin: 0.8em 0 0.4em; }}
  h2 {{ font-size: 16pt; font-weight: bold; margin: 0.7em 0 0.3em; }}
  h3 {{ font-size: 13pt; font-weight: bold; margin: 0.6em 0 0.3em; }}
  h4, h5, h6 {{ font-size: 11pt; font-weight: bold; margin: 0.5em 0 0.2em; }}
  p  {{ margin: 0 0 0.5em; }}
  b, strong {{ font-weight: bold; }}
  i, em {{ font-style: italic; }}
  u {{ text-decoration: underline; }}
  ul, ol {{ margin: 0.4em 0 0.4em 1.5em; padding: 0; }}
  li {{ margin-bottom: 0.2em; }}
  table {{ border-collapse: collapse; width: 100%; margin: 0.8em 0; }}
  td, th {{ border: 1px solid #999; padding: 4px 8px; vertical-align: top; }}
  th {{ background: #f0f0f0; font-weight: bold; }}
  img {{ max-width: 100%; height: auto; }}
</style>
</head>
<body>{result.value}</body>
</html>"""

            weasyprint.HTML(string=html).write_pdf(str(pdf_path))
            return pdf_path.exists()
        except Exception:
            return False

    def _try_python_docx_reportlab() -> bool:
        # Last-resort fallback: python-docx + reportlab only — no system libs needed.
        # Preserves bold/italic/headings/tables but drops images (needs pango for those).
        try:
            import io as _io
            from docx import Document as _Document
            from docx.oxml.ns import qn as _qn
            from docx.text.paragraph import Paragraph as _DocxPara
            from docx.table import Table as _DocxTable
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            from reportlab.lib import colors as _colors
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                Image as _RLImage,
                Paragraph as _RLPara,
                SimpleDocTemplate,
                Spacer,
                Table as _RLTable,
                TableStyle,
            )

            _doc = _Document(str(docx_path))

            _S = {
                "n":  ParagraphStyle("rln",  fontName="Helvetica",      fontSize=11, leading=16, spaceAfter=4),
                "h1": ParagraphStyle("rlh1", fontName="Helvetica-Bold", fontSize=20, leading=26, spaceBefore=10, spaceAfter=6),
                "h2": ParagraphStyle("rlh2", fontName="Helvetica-Bold", fontSize=16, leading=22, spaceBefore=8,  spaceAfter=4),
                "h3": ParagraphStyle("rlh3", fontName="Helvetica-Bold", fontSize=13, leading=18, spaceBefore=6,  spaceAfter=4),
            }
            _ALIGN = {
                WD_ALIGN_PARAGRAPH.CENTER:  TA_CENTER,
                WD_ALIGN_PARAGRAPH.RIGHT:   TA_RIGHT,
                WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY,
            }

            def _style(base, alignment):
                ta = _ALIGN.get(alignment)
                if ta is None:
                    return base
                return ParagraphStyle(base.name + "_a", parent=base, alignment=ta)

            def _runs_to_markup(para):
                parts = []
                for run in para.runs:
                    t = (run.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if not t:
                        continue
                    if run.bold:      t = f"<b>{t}</b>"
                    if run.italic:    t = f"<i>{t}</i>"
                    if run.underline: t = f"<u>{t}</u>"
                    parts.append(t)
                return "".join(parts)

            def _images_in_para(para):
                imgs = []
                for run in para.runs:
                    for blip in run._r.findall(".//" + _qn("a:blip")):
                        rid = blip.get(_qn("r:embed"))
                        if rid and rid in _doc.part.rels:
                            rel = _doc.part.rels[rid]
                            if "image" in rel.reltype:
                                imgs.append(rel.target_part.blob)
                return imgs

            story = []

            def _add_para(para):
                for img_blob in _images_in_para(para):
                    try:
                        from PIL import Image as _PILImg
                        pil = _PILImg.open(_io.BytesIO(img_blob))
                        w, h = pil.size
                        max_w, max_h = 5.5 * inch, 2.5 * inch
                        scale = min(max_w / w, max_h / h, 1.0)
                        story.append(_RLImage(_io.BytesIO(img_blob), width=w * scale, height=h * scale))
                        story.append(Spacer(1, 4))
                    except Exception:
                        pass

                markup = _runs_to_markup(para)
                if not markup.strip():
                    story.append(Spacer(1, 4))
                    return

                sname = (para.style.name or "").lower()
                if "heading 1" in sname:   base = _S["h1"]
                elif "heading 2" in sname: base = _S["h2"]
                elif "heading 3" in sname: base = _S["h3"]
                else:                      base = _S["n"]

                st = _style(base, para.alignment)
                try:
                    story.append(_RLPara(markup, st))
                except Exception:
                    story.append(_RLPara((para.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), base))

            def _add_table(table):
                data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        ct = "".join(
                            (r.text or "") for p in cell.paragraphs for r in p.runs
                        ).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        row_data.append(_RLPara(ct, _S["n"]))
                    data.append(row_data)
                if data:
                    t = _RLTable(data, repeatRows=1, hAlign="LEFT")
                    t.setStyle(TableStyle([
                        ("GRID",         (0, 0), (-1, -1), 0.5, _colors.grey),
                        ("BACKGROUND",   (0, 0), (-1,  0), _colors.lightgrey),
                        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
                        ("TOPPADDING",   (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 8))

            for child in _doc.element.body:
                tag = child.tag
                if tag == _qn("w:p"):
                    _add_para(_DocxPara(child, _doc))
                elif tag == _qn("w:tbl"):
                    _add_table(_DocxTable(child, _doc))

            if not story:
                story.append(_RLPara("(empty document)", _S["n"]))

            SimpleDocTemplate(
                str(pdf_path), pagesize=letter,
                leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
            ).build(story)
            return pdf_path.exists()
        except Exception:
            return False

    ok = False
    if sys.platform == "win32":
        ok = _try_docx2pdf() or _try_soffice() or _try_mammoth_weasyprint() or _try_python_docx_reportlab()
    else:
        ok = _try_soffice() or _try_docx2pdf() or _try_mammoth_weasyprint() or _try_python_docx_reportlab()

    if ok:
        return pdf_path

    raise RuntimeError(
        "Could not convert the Word (.docx) file to PDF.\n"
        "For best results install LibreOffice: brew install --cask libreoffice\n"
        "Or install pango for image support: brew install pango"
    )


def merge_pdfs(pdf_paths: List[Path], output_path: Path) -> None:
    """
    Merge PDFs in the given order into a single file.
    """
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    writer = PdfWriter()
    for p in pdf_paths:
        p = Path(p).resolve()
        if not p.exists():
            raise FileNotFoundError(p)
        reader = PdfReader(str(p))
        for page in reader.pages:
            writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)


def add_numbered_header(input_path: Path, output_path: Path) -> None:
    """
    Add "Pag. n/total" to each page (header). Uses the existing add_page_numbers module.
    """
    from add_page_numbers import add_numbers_to_pdf
    add_numbers_to_pdf(Path(input_path), Path(output_path))


def build_merged_pdf(
    file_paths: List[Path],
    output_path: Path,
    enumerate: bool = False,
    temp_dir: Path | None = None,
) -> Path:
    """
    Main pipeline: convert any .docx to PDF, merge all in order, optionally add page numbers.

    file_paths: List of paths to .pdf or .docx files (order preserved).
    output_path: Where to write the final PDF.
    enumerate: If True, run add_numbered_header on the merged document before saving.
    temp_dir: Optional directory for intermediate files; uses tempfile if not set.

    Returns the path to the final PDF.
    """
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    use_temp = temp_dir is None
    if use_temp:
        temp_dir = Path(tempfile.mkdtemp())
    else:
        temp_dir = Path(temp_dir).resolve()
        temp_dir.mkdir(parents=True, exist_ok=True)

    try:
        pdf_paths: List[Path] = []
        for p in file_paths:
            p = Path(p).resolve()
            if not p.exists():
                raise FileNotFoundError(p)
            if p.suffix.lower() == ".docx":
                pdf_paths.append(convert_docx_to_pdf(p, output_dir=temp_dir))
            elif p.suffix.lower() == ".pdf":
                pdf_paths.append(p)
            else:
                raise ValueError(f"Unsupported format: {p} (use .pdf or .docx)")

        if not pdf_paths:
            raise ValueError("No PDF or DOCX files to merge.")

        merged_path = temp_dir / "merged.pdf"
        merge_pdfs(pdf_paths, merged_path)

        if enumerate:
            add_numbered_header(merged_path, output_path)
        else:
            shutil.copy2(merged_path, output_path)

        return output_path
    finally:
        if use_temp and temp_dir.exists():
            shutil.rmtree(temp_dir, ignore_errors=True)
